import os
from flask import Flask, url_for, render_template, flash, redirect, request, send_from_directory
from dotenv import load_dotenv
from webforms import LoginForm, RegistrationForm
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.exc import IntegrityError
from datetime import datetime
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from flask_login import UserMixin, login_user, login_required, logout_user, current_user, LoginManager
from flask_migrate import Migrate
from flask_wtf.csrf import CSRFProtect
from docx import Document
from openai_functions import translation_openai

load_dotenv()

# Define the path to the upload folder
UPLOAD_FOLDER = 'user-files'
TRANSLATED_FILES_FOLDER = 'translated-files'
ALLOWED_EXTENSIONS = {'docx'}

csrf = CSRFProtect()
app = Flask(__name__)
csrf.init_app(app)
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DB_URI')
app.config['SECRET_KEY'] = os.environ.get('FORM_KEY')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

db = SQLAlchemy(app)
# migrate changes of a schema to the db
migrate = Migrate(app, db)

login_manager = LoginManager()
login_manager.init_app(app)

login_manager.login_view = 'login'


@login_manager.user_loader
def load_user(user_id):
    return Users.query.get(int(user_id))


# login page
@app.route("/", methods=['GET', 'POST'])
def login():
    form = LoginForm()
    email = None
    password = None
    errors = None

    if form.validate_on_submit():
        email = form.email.data
        password = form.password.data

        try:
            existing_user = Users.query.filter_by(email=email).first()
            if existing_user and check_password_hash(existing_user.password_hash, password):
                login_user(existing_user)
                return redirect(url_for('dashboard'))
            else:
                flash("Invalid email or password. Please try again.",
                      category='error')

        except IntegrityError:
            # Handle database integrity errors (e.g., unique constraint violation)
            db.session.rollback()  # Rollback the transaction
            print("An error occurred during registration. Please try again.", 'error')
    else:
        errors = form.errors

    return render_template('auth/login.html',
                           form=form,
                           email=email,
                           password=password,
                           errors=errors
                           )


# logout
@app.route("/logout", methods=['GET', 'POST'])
@login_required
def logout():
    logout_user()
    flash("You have been logged out!", category='success')
    return redirect(url_for('login'))


@app.route("/register", methods=['GET', 'POST'])
def register():
    form = RegistrationForm()
    email = None
    password = None
    errors = None

    users = Users.query.all()

    if form.validate_on_submit():
        email = form.email.data
        password = form.password.data
        selected_role = form.role.data

        try:
            existing_user = Users.query.filter_by(email=email).first()
            if existing_user:
                flash("Email already in use. Please choose another.",
                      category='error')

            else:
                hashed_password = generate_password_hash(password)
                new_user = Users(
                    email=email, password_hash=hashed_password, role=selected_role)
                db.session.add(new_user)
                db.session.commit()
                flash("Registration successful", category='success')
                return redirect(url_for('register'))

        except IntegrityError:
            # Handle database integrity errors (e.g., unique constraint violation)
            db.session.rollback()  # Rollback the transaction
            flash(
                "An error occurred during registration. Please try again.",  category='error')
    else:
        errors = form.errors

    return render_template('auth/register.html', form=form, email=email, password=password, errors=errors, users=users)


@app.route("/delete_user/<int:user_id>", methods=['POST'])
@login_required
def delete_user(user_id):
    #check if the user is an admin
    if current_user.role != 'admin':
        flash("You are not authorized to delete users", category='error')
        return redirect(url_for('register'))
    try:
        # get user from db
        user = Users.query.get(user_id)

        # check if user exists
        if user:
            # delete the user from db
            db.session.delete(user)
            db.session.commit()
            flash("user deleted", category='error')
        else:
            flash("something went wrong")
    except IntegrityError:
        db.session.rollback()  # Rollback the transaction
        flash("something went wrong when deleting the user", category='error')

    return redirect(url_for('register'))


@app.route("/delete_file/<int:file_id>", methods=['POST'])
@login_required
def delete_file(file_id):
    try:
        # Get the file from both UploadedFiles and TranslatedFiles tables
        uploaded_file = UploadedFiles.query.get(file_id)
        translated_file = TranslatedFiles.query.get(file_id)

        # Check if the file exists and belongs to the user in either table
        if uploaded_file and uploaded_file.user_id == current_user.id:
            # Delete the file from the server
            if os.path.exists(uploaded_file.file_path):
                os.remove(uploaded_file.file_path)

            # Delete the file from the UploadedFiles table
            db.session.delete(uploaded_file)
            db.session.commit()
            flash("File deleted", category='warning')

            # Check if the folder is empty and delete it
            folder_path = os.path.dirname(uploaded_file.file_path)
            if not os.listdir(folder_path):
                os.rmdir(folder_path)
        elif translated_file and translated_file.user_id == current_user.id:
            # Delete the file from the server
            if os.path.exists(translated_file.file_path):
                os.remove(translated_file.file_path)

            # Delete the file from the TranslatedFiles table
            db.session.delete(translated_file)
            db.session.commit()
            flash("File deleted", category='warning')

            # Check if the folder is empty and delete it
            folder_path = os.path.dirname(translated_file.file_path)
            if not os.listdir(folder_path):
                os.rmdir(folder_path)
        else:
            flash("Something went wrong")

    except IntegrityError:
        db.session.rollback()
        flash("Something went wrong when deleting the file", category='error')

    return redirect(url_for('translation'))





# dashboard page
@app.route("/dashboard")
@login_required
def dashboard():
    is_admin = False 

    if current_user.role == 'admin':
        is_admin = True

    return render_template('dashboard.html', is_admin=is_admin)


@app.route("/download_file/<int:file_id>", methods=['GET'])
@login_required
def download_file(file_id):
    try:
        # Fetch the file from the database based on file_id
        file = TranslatedFiles.query.get(file_id)

        if file and file.user_id == current_user.id:
            # Download the file from the server
            return send_from_directory(os.path.dirname(file.file_path), file.file_name, as_attachment=True)

        else:
            flash("File not found or you don't have permission to download it", 'error')
    except Exception as e:
        # Handle any exceptions that may occur during download
        print(f"Error downloading file: {e}")
        flash("An error occurred during download", 'succes')

    # Redirect to a relevant URL after processing
    return redirect(url_for('translation'))


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route("/translate_files", methods=['POST'])
@login_required
def translate_files():
    try:
        # Get the selected file for translation
        selected_file_id = request.form.get('selected_file')

        # Get the selected languages for translation
        language_from = request.form.getlist('languageToTranslatefrom')
        language_to = request.form.getlist('languagesToTranslateTo')  # If multiple languages are selected, use getlist

        print(f'Selected File ID: {selected_file_id}')
        print(f'Language From: {language_from}')
        print(f'Languages To: {language_to}')

        # Fetch the file from the database based on selected_file_id
        file = UploadedFiles.query.get(selected_file_id)

        if file:
            # Use python-docx to read the contents of the .docx file
            doc = Document(os.path.join(file.file_path))
            file_contents = '\n'.join([para.text for para in doc.paragraphs])

            # Translate the text using OpenAI
            translated_text = translation_openai(file_contents, language_from, language_to)
            print(f'Translated Text: {translated_text}')

            # Define a folder path for translated files
            translated_folder_path = os.path.join(app.config['UPLOAD_FOLDER'], TRANSLATED_FILES_FOLDER, str(current_user.id))
            os.makedirs(translated_folder_path, exist_ok=True)

            # Construct the translated file name based on the original filename and target language
            original_filename = file.file_name
            target_language = language_to[0]  # Assuming only one target language is selected
            translated_file_name = f"{original_filename.split('.')[0]}_{target_language}.docx"

            # Define the file path for the translated file
            translated_file_path = os.path.join(translated_folder_path, translated_file_name)

            # Create a new .docx document
            doc = Document()
            doc.add_paragraph(translated_text)

            # Save the document
            doc.save(translated_file_path)

            if translated_file_name and translated_file_path:
                # Create a new record in the TranslatedFiles table
                new_translated_file = TranslatedFiles(
                    file_name=translated_file_name,
                    file_path=translated_file_path,
                    user_id=current_user.id
                )
                db.session.add(new_translated_file)
                db.session.commit()
                flash("File translated and saved successfully", category='success')

    except Exception as e:
        # Handle any exceptions that may occur during translation
        print(f"Error translating file: {e}")

    # Redirect to a relevant URL after processing
    return redirect(url_for('translation'))

# translation page
@app.route("/translation", methods=['GET', 'POST'])
@login_required
def translation():
    languages_to_translate_from = ['Norwegian', 'English', 'Spanish', 'French', 'German', 'Chinese (Simplified)', 'Chinese (Traditional)', 'Japanese', 'Korean', 'Russian', 'Arabic']
    languages_to_translate_to = ["English", "Spanish", "French", "German",
    "Chinese (Simplified)", "Chinese (Traditional)", "Japanese", "Korean", "Russian", "Arabic", "Italian", "Portuguese", "Dutch", "Swedish", "Norwegian", "Danish", "Finnish", "Greek", "Turkish", "Hindi"
]

    if request.method == 'POST':
        # check if the post request has the file part
        if 'files' not in request.files:
            print('No file part')
            return redirect(request.url)
        files = request.files.getlist('files')

        try:
            for file in files:
                if file.filename == '':
                    flash('No selected file', 'error')
                    continue

                if file and allowed_file(file.filename):
                    filename = secure_filename(file.filename)
                    user_upload_folder = os.path.join(app.config['UPLOAD_FOLDER'], str(current_user.id))

                    #ensure the user's upload folder exists
                    os.makedirs(user_upload_folder, exist_ok=True)

                    file.save(os.path.join(user_upload_folder, filename))

                    # Create and add a new UploadedFiles record to the database
                    new_file = UploadedFiles(
                        file_name=filename,
                        file_path=os.path.join(user_upload_folder, filename),
                        user_id=current_user.id  # Assign the current user's ID
                    )
                    db.session.add(new_file)
                    db.session.commit()  # Commit changes to the database
                    flash('File(s) uploaded successfully', category='success')

                else:
                    flash('Invalid file type. Only .docx files are allowed', 'error')
        except Exception as e:
            db.session.rollback()  # Rollback the transaction in case of an exception
            print(f'An error occurred: {str(e)}', 'error')

    # Fetch the user's files from the db
    user_files = UploadedFiles.query.filter_by(user_id=current_user.id).all()
    user_translated_files = TranslatedFiles.query.filter_by(user_id=current_user.id).all()

    return render_template('translation.html', user_files=user_files, user_translated_files=user_translated_files, languages_to_translate_to=languages_to_translate_to, languages_to_translate_from=languages_to_translate_from)



# invalid url
@app.errorhandler(404)
def page_not_found(error):
    return render_template('errors/404.html'), 404


# internal server error
@app.errorhandler(500)
def special_exception_handler(error):
    return render_template('errors/500.html'), 500


#models
class Users(db.Model, UserMixin):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(120), nullable=False, unique=True)
    password_hash = db.Column(db.String(128), nullable=False)
    role = db.Column(db.String(120), default='user')
    date_added = db.Column(db.DateTime, default=datetime.utcnow)
    files = db.relationship('UploadedFiles', backref='user', lazy=True)
    translated_files = db.relationship('TranslatedFiles', backref='user', lazy=True)

    @property
    def password(self):
        raise AttributeError('password is not a readable attribute!')

    @password.setter
    def password(self, password):
        self.password_hash = generate_password_hash(password)

    def verify_password(self, password):
        return check_password_hash(self.password_hash, password)

    # create a string
    def __repr__(self):
        return f'<User {self.email}>'

class UploadedFiles(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    file_name = db.Column(db.String(120), nullable=False)
    file_path = db.Column(db.String(120), nullable=False)
    date_added = db.Column(db.DateTime, default=datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)


    # create a string
    def __repr__(self):
        return f'<UploadedFiles {self.file_name}>'

class TranslatedFiles(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    file_name = db.Column(db.String(120), nullable=False)
    file_path = db.Column(db.String(120), nullable=False)
    date_added = db.Column(db.DateTime, default=datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)

    #creatre a string
    def __repr__(self):
        return f'<TranslatedFiles {self.file_name}>'